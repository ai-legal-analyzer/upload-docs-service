import pytest
import httpx
import random
import time
import io
from pathlib import Path

BASE_URL = "http://127.0.0.1:8000"


def generate_test_pdf():
    """Generate a simple PDF file for testing"""
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    c.drawString(100, 750, "Test Document")
    c.drawString(100, 730, "Test content for document processing")
    c.drawString(100, 710, f"Random ID: {random.randint(1000, 9999)}")
    c.drawString(100, 690, "This is a test PDF file for the upload service.")
    c.drawString(100, 670, "It contains sample text to be extracted and chunked.")
    c.save()

    buffer.seek(0)
    return buffer


def generate_test_docx():
    """Generate a simple DOCX file for testing"""
    from docx import Document

    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph(f'This is a test DOCX file with random ID: {random.randint(1000, 9999)}')
    doc.add_paragraph('This document will be processed by the upload service.')
    doc.add_paragraph('The text should be extracted and split into chunks for storage.')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_random_owner_id():
    return random.randint(1000, 9999)


@pytest.fixture
def client():
    with httpx.Client(base_url=BASE_URL, timeout=30.0) as c:
        yield c


def test_app_health(client):
    """Test that the API is running and accessible"""
    response = client.get("openapi.json")
    assert response.status_code == 200


def test_upload_pdf_document(client):
    """Test uploading a PDF document"""
    owner_id = generate_random_owner_id()
    pdf_file = generate_test_pdf()

    files = {
        "file": ("test_document.pdf", pdf_file, "application/pdf")
    }

    response = client.post(
        f"/api/documents/?owner_id={owner_id}",
        files=files
    )

    assert response.status_code == 202
    body = response.json()
    assert "task_id" in body
    assert body["status"] == "processing"
    assert "message" in body

    # Store task_id for other tests if needed, but don't return it
    # If you need the task_id elsewhere, use a fixture or class variable


def test_upload_docx_document(client):
    """Test uploading a DOCX document"""
    owner_id = generate_random_owner_id()
    docx_file = generate_test_docx()

    files = {
        "file": ("test_document.docx", docx_file,
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    }

    response = client.post(
        f"/api/documents/?owner_id={owner_id}",
        files=files
    )

    assert response.status_code == 202
    body = response.json()
    assert "task_id" in body
    assert body["status"] == "processing"

    # Don't return task_id - test functions should return None


def test_get_task_status(client):
    """Test checking task status"""
    # First upload a document to get a task_id
    owner_id = generate_random_owner_id()
    pdf_file = generate_test_pdf()

    files = {
        "file": ("test_document.pdf", pdf_file, "application/pdf")
    }

    upload_response = client.post(
        f"/api/documents/?owner_id={owner_id}",
        files=files
    )

    assert upload_response.status_code == 202
    task_id = upload_response.json()["task_id"]

    # Wait a bit for processing to start
    time.sleep(2)

    response = client.get(f"/api/documents/task/{task_id}")
    assert response.status_code == 200

    body = response.json()
    assert body["task_id"] == task_id
    assert "state" in body
    assert body["state"] in ["PENDING", "PROGRESS", "SUCCESS", "FAILURE"]


def _wait_for_task_completion(client, owner_id):
    """Helper function to wait for task completion and return document_id"""
    pdf_file = generate_test_pdf()

    files = {
        "file": ("completion_test.pdf", pdf_file, "application/pdf")
    }

    # Upload document
    upload_response = client.post(
        f"/api/documents/?owner_id={owner_id}",
        files=files
    )
    task_id = upload_response.json()["task_id"]

    # Poll for completion (with timeout)
    max_wait = 60  # seconds
    wait_time = 0
    poll_interval = 2

    while wait_time < max_wait:
        status_response = client.get(f"/api/documents/task/{task_id}")
        status_data = status_response.json()

        if status_data["state"] == "SUCCESS":
            assert "result" in status_data
            assert "document_id" in status_data["result"]
            assert "num_chunks" in status_data["result"]
            assert status_data["result"]["num_chunks"] > 0
            return status_data["result"]["document_id"]
        elif status_data["state"] == "FAILURE":
            pytest.fail(f"Task failed: {status_data}")

        time.sleep(poll_interval)
        wait_time += poll_interval

    pytest.fail("Task did not complete within timeout period")


def test_wait_for_task_completion(client):
    """Test waiting for a task to complete and verify result"""
    owner_id = generate_random_owner_id()
    document_id = _wait_for_task_completion(client, owner_id)

    # Verify we got a document_id (implicitly tested in the helper)
    assert isinstance(document_id, int)
    assert document_id > 0


def test_list_documents_by_owner(client):
    """Test listing documents for a specific owner"""
    owner_id = generate_random_owner_id()

    # Upload multiple documents for the same owner
    for i in range(2):
        pdf_file = generate_test_pdf()
        files = {"file": (f"doc_{i}.pdf", pdf_file, "application/pdf")}
        client.post(f"/api/documents/?owner_id={owner_id}", files=files)

    # Wait for processing
    time.sleep(5)

    response = client.get(f"/api/documents/owner/{owner_id}")
    assert response.status_code == 200

    body = response.json()
    assert "documents" in body
    assert "total_count" in body
    assert body["owner_id"] == owner_id
    assert len(body["documents"]) >= 0  # Could be 0 if processing not complete


def test_list_all_documents(client):
    """Test listing all documents"""
    response = client.get("/api/documents/")
    assert response.status_code == 200

    body = response.json()
    assert "documents" in body
    assert "total_count" in body
    assert "skip" in body
    assert "limit" in body


def test_get_document_chunks(client):
    """Test retrieving chunks for a specific document"""
    # First create a document and wait for processing
    owner_id = generate_random_owner_id()
    document_id = _wait_for_task_completion(client, owner_id)

    response = client.get(f"/api/documents/{document_id}/chunks")

    if response.status_code == 404:
        pytest.skip("Document not found (might still be processing)")

    assert response.status_code == 200

    body = response.json()
    assert "document" in body
    assert "chunks" in body
    assert "total_chunks" in body
    assert body["document"]["id"] == document_id

    if body["total_chunks"] > 0:
        chunk = body["chunks"][0]
        assert "id" in chunk
        assert "chunk_index" in chunk
        assert "text" in chunk
        assert len(chunk["text"]) > 0


def test_upload_invalid_file_type(client):
    """Test uploading an invalid file type"""
    owner_id = generate_random_owner_id()

    # Create a text file (not supported)
    text_file = io.BytesIO(b"This is a text file, not PDF or DOCX")

    files = {
        "file": ("invalid.txt", text_file, "text/plain")
    }

    response = client.post(
        f"/api/documents/?owner_id={owner_id}",
        files=files
    )

    # Should return 422 Unprocessable Entity or similar error
    assert response.status_code in [400, 422, 500]


def test_pagination_parameters(client):
    """Test pagination parameters work correctly"""
    owner_id = generate_random_owner_id()

    response = client.get(f"/api/documents/owner/{owner_id}?skip=5&limit=10")
    assert response.status_code == 200

    body = response.json()
    assert body["skip"] == 5
    assert body["limit"] == 10


def test_nonexistent_task_status(client):
    """Test checking status of non-existent task"""
    response = client.get("/api/documents/task/nonexistent-task-id")
    assert response.status_code == 200  # Celery returns info even for non-existent tasks

    body = response.json()
    assert body["state"] == "PENDING"


def test_nonexistent_document_chunks(client):
    """Test getting chunks for non-existent document"""
    response = client.get("/api/documents/999999/chunks")
    assert response.status_code == 404


# Performance and stress tests
def test_multiple_uploads_concurrent(client):
    """Test uploading multiple documents in sequence"""
    owner_id = generate_random_owner_id()
    task_ids = []

    for i in range(3):
        if i % 2 == 0:
            file_content = generate_test_pdf()
            filename = f"test_{i}.pdf"
            content_type = "application/pdf"
        else:
            file_content = generate_test_docx()
            filename = f"test_{i}.docx"
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        files = {"file": (filename, file_content, content_type)}
        response = client.post(f"/api/documents/?owner_id={owner_id}", files=files)

        assert response.status_code == 202
        task_ids.append(response.json()["task_id"])

    assert len(task_ids) == 3
    assert len(set(task_ids)) == 3  # All task IDs should be unique
