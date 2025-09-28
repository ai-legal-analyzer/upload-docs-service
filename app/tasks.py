import tempfile
from pathlib import Path
from typing import Optional, Any
from celery import current_task
from sqlalchemy.ext.asyncio import AsyncSession, create_async_engine
from sqlalchemy.orm import sessionmaker

from app.celery_app import celery_app
from app.models.document import Document
from app.models.chunk import DocumentChunk
from app.backend.db import DATABASE_URL

# Create async engine for tasks
# engine = create_async_engine(DATABASE_URL)
# AsyncSessionLocal = sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)

ALLOWED_EXTENSIONS = {"pdf", "docx"}
CHUNK_SIZE = 3000  # characters


def parse_pdf(file_path: Path) -> str:
    """Extract text from PDF file."""
    from PyPDF2 import PdfReader
    reader = PdfReader(file_path)
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def parse_docx(file_path: Path) -> str:
    """Extract text from DOCX file."""
    from docx import Document as DocxDocument
    doc = DocxDocument(file_path)
    return "\n".join(para.text for para in doc.paragraphs)


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE) -> list[str]:
    """Split text into chunks of specified size."""
    return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]


@celery_app.task(bind=True)
def process_document(
    self,
    file_content: bytes,
    filename: str,
    content_type: str,
    owner_id: int
) -> dict[str, Any]:
    """
    Background task to process uploaded documents.
    
    Args:
        file_content: The file content as bytes
        filename: Original filename
        content_type: MIME type of the file
        owner_id: ID of the document owner
        
    Returns:
        dict with document_id and num_chunks
    """
    try:
        # Update task state
        self.update_state(
            state="PROGRESS",
            meta={"status": "Processing document", "progress": 10}
        )
        
        # Create temporary file
        ext = Path(filename).suffix
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(file_content)
            tmp_path = Path(tmp.name)
        
        try:
            # Extract text based on file type
            self.update_state(
                state="PROGRESS",
                meta={"status": "Extracting text", "progress": 30}
            )
            
            match tmp_path.suffix.lower():
                case ".pdf":
                    text = parse_pdf(tmp_path)
                case ".docx":
                    text = parse_docx(tmp_path)
                case _:
                    raise ValueError("Unsupported file type")
            
            if not text.strip():
                raise ValueError("No text extracted from file")
            
            # Chunk the text
            self.update_state(
                state="PROGRESS",
                meta={"status": "Chunking text", "progress": 60}
            )
            
            chunks = chunk_text(text)
            
            # Save to database
            self.update_state(
                state="PROGRESS",
                meta={"status": "Saving to database", "progress": 80}
            )
            
            # Use async context to save to database
            import asyncio
            
            async def save_document():
                from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession
                from sqlalchemy.orm import sessionmaker

                engine = create_async_engine(DATABASE_URL)
                AsyncSessionLocal = sessionmaker(engine, class_=AsyncSession, expire_on_commit=False)
                async with AsyncSessionLocal() as db:
                    async with db.begin():
                        doc = Document(
                            filename=filename,
                            content_type=content_type,
                            num_chunks=len(chunks),
                            owner_id=owner_id
                        )
                        db.add(doc)
                        await db.flush()  # Get the ID before commit
                        
                        db.add_all(
                            DocumentChunk(
                                document_id=doc.id,
                                chunk_index=idx,
                                text=chunk
                            )
                            for idx, chunk in enumerate(chunks)
                        )
                        
                        return doc.id, len(chunks)
            
            # Run the async function
            doc_id, num_chunks = asyncio.run(save_document())
            
            # Update final state
            self.update_state(
                state="SUCCESS",
                meta={
                    "status": "Document processed successfully",
                    "progress": 100,
                    "document_id": doc_id,
                    "num_chunks": num_chunks
                }
            )
            
            return {
                "document_id": doc_id,
                "num_chunks": num_chunks,
                "status": "success"
            }
            
        finally:
            # Clean up temporary file
            tmp_path.unlink(missing_ok=True)
            
    except Exception as exc:
        raise


# @celery_app.task
# def cleanup_old_documents(days_old: int = 30) -> dict[str, any]:
#     """
#     Background task to clean up old documents.
#
#     Args:
#         days_old: Number of days after which documents should be deleted
#
#     Returns:
#         dict with number of deleted documents
#     """
#     try:
#         import asyncio
#         from datetime import datetime, timedelta
#         from sqlalchemy import delete
#
#         async def delete_old_documents():
#             async with AsyncSessionLocal() as db:
#                 cutoff_date = datetime.utcnow() - timedelta(days=days_old)
#
#                 # Delete documents older than cutoff_date
#                 result = await db.execute(
#                     delete(Document).where(Document.upload_time < cutoff_date)
#                 )
#                 await db.commit()
#
#                 return result.rowcount
#
#         # Run the async function
#         loop = asyncio.new_event_loop()
#         asyncio.set_event_loop(loop)
#         try:
#             deleted_count = loop.run_until_complete(delete_old_documents())
#         finally:
#             loop.close()
#
#         return {
#             "deleted_documents": deleted_count,
#             "status": "success"
#         }
#
#     except Exception as exc:
#         return {
#             "error": str(exc),
#             "status": "error"
#         }